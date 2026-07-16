#!/usr/bin/env python3
"""
Chenier Environmental Consulting
Site Location Map Generator — Streamlit Web App
Deploy to Streamlit Community Cloud (free).
"""

import io, math, zipfile
from datetime import datetime

import streamlit as st
import requests
from PIL import Image, ImageDraw
from shapely.geometry import Polygon
from shapely.ops import unary_union
from lxml import etree
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Constants ──────────────────────────────────────────────────────────────────
OSM_TILE_URL  = 'https://tile.openstreetmap.org/{z}/{x}/{y}.png'
OSM_HEADERS   = {'User-Agent': 'ChenierLocator/1.0 (environmental consulting)'}
TILE_SIZE     = 256
MIN_BUFFER_DEG = 0.004
SITE_COLOR    = (255, 0, 0)
IMG_W_PX      = 1125
IMG_H_PX      = 1275
JPEG_QUALITY  = 85
PAGE_W        = Inches(8.5)
PAGE_H        = Inches(11.0)
MAR_L         = Inches(0.7)
MAR_R         = Inches(0.25)
MAR_TOP       = Inches(0.5)
MAR_BOT       = Inches(0.5)
MAP_IMG_W     = Inches(7.5)
MAP_IMG_H     = Inches(8.5)

# north_arrow.jpeg should live in the same repo as this app
NORTH_ARROW_PATH = 'north_arrow.jpeg'


# ── KMZ / KML parsing ─────────────────────────────────────────────────────────
def parse_kmz_bytes(file_bytes, filename):
    """Parse KMZ/KML from uploaded bytes. Returns (geom, bounds)."""
    if filename.lower().endswith('.kmz'):
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
            names = [n for n in z.namelist() if n.lower().endswith('.kml')]
            if not names:
                raise ValueError("No KML found inside KMZ")
            kml = z.read('doc.kml' if 'doc.kml' in names else names[0])
    else:
        kml = file_bytes

    root  = etree.fromstring(kml)
    polys = []
    for el in root.iter('{http://www.opengis.net/kml/2.2}coordinates'):
        pts = []
        for tok in el.text.strip().split():
            p = tok.split(',')
            if len(p) >= 2:
                try:
                    pts.append((float(p[0]), float(p[1])))
                except:
                    pass
        if len(pts) >= 3:
            polys.append(Polygon(pts))

    if not polys:
        raise ValueError("No polygon found in KMZ/KML")
    geom = unary_union(polys)
    return geom, geom.bounds


# ── OSM tile math ─────────────────────────────────────────────────────────────
def deg_to_tile(lat_deg, lon_deg, zoom):
    lat_r = math.radians(lat_deg)
    n     = 2 ** zoom
    x     = (lon_deg + 180.0) / 360.0 * n
    y     = (1.0 - math.log(math.tan(lat_r) + 1.0 / math.cos(lat_r)) / math.pi) / 2.0 * n
    return x, y


def tile_to_deg(x, y, zoom):
    n     = 2 ** zoom
    lon   = x / n * 360.0 - 180.0
    lat_r = math.atan(math.sinh(math.pi * (1 - 2 * y / n)))
    lat   = math.degrees(lat_r)
    return lat, lon


def choose_zoom(bounds, target_px_w, target_px_h):
    min_lng, min_lat, max_lng, max_lat = bounds
    for z in range(17, 8, -1):
        tx0, ty0 = deg_to_tile(max_lat, min_lng, z)
        tx1, ty1 = deg_to_tile(min_lat, max_lng, z)
        span_px_x = (tx1 - tx0) * TILE_SIZE
        span_px_y = (ty1 - ty0) * TILE_SIZE
        if span_px_x <= target_px_w * 0.75 and span_px_y <= target_px_h * 0.75:
            return z
    return 10


def fetch_tile(z, x, y, cache):
    key = (z, x, y)
    if key in cache:
        return cache[key]
    url = OSM_TILE_URL.format(z=z, x=x, y=y)
    try:
        r = requests.get(url, headers=OSM_HEADERS, timeout=15)
        r.raise_for_status()
        img = Image.open(io.BytesIO(r.content)).convert('RGB')
        cache[key] = img
        return img
    except Exception:
        blank = Image.new('RGB', (TILE_SIZE, TILE_SIZE), (200, 200, 200))
        cache[key] = blank
        return blank


def build_osm_canvas(bounds, zoom, out_w, out_h, progress=None):
    min_lng, min_lat, max_lng, max_lat = bounds

    tx0f, ty0f = deg_to_tile(max_lat, min_lng, zoom)
    tx1f, ty1f = deg_to_tile(min_lat, max_lng, zoom)
    tx0, ty0 = int(tx0f), int(ty0f)
    tx1, ty1 = int(tx1f), int(ty1f)

    cols = tx1 - tx0 + 1
    rows = ty1 - ty0 + 1
    canvas_w = cols * TILE_SIZE
    canvas_h = rows * TILE_SIZE
    canvas = Image.new('RGB', (canvas_w, canvas_h))

    cache = {}
    total = cols * rows
    done  = 0
    for ty in range(ty0, ty1 + 1):
        for tx in range(tx0, tx1 + 1):
            tile = fetch_tile(zoom, tx, ty, cache)
            canvas.paste(tile, ((tx - tx0) * TILE_SIZE, (ty - ty0) * TILE_SIZE))
            done += 1
            if progress:
                progress(done / total)

    nw_lat, nw_lng = tile_to_deg(tx0,     ty0,     zoom)
    se_lat, se_lng = tile_to_deg(tx1 + 1, ty1 + 1, zoom)

    # Pad bounds 5% before crop
    lng_span = max_lng - min_lng
    lat_span = max_lat - min_lat
    pad_lng  = lng_span * 0.05
    pad_lat  = lat_span * 0.05
    crop_min_lng = max(nw_lng, min_lng - pad_lng)
    crop_max_lng = min(se_lng, max_lng + pad_lng)
    crop_max_lat = min(nw_lat, max_lat + pad_lat)
    crop_min_lat = max(se_lat, min_lat - pad_lat)

    def geo_to_canvas_px(lng, lat):
        tx_f, ty_f = deg_to_tile(lat, lng, zoom)
        return (tx_f - tx0) * TILE_SIZE, (ty_f - ty0) * TILE_SIZE

    cx0, cy0 = geo_to_canvas_px(crop_min_lng, crop_max_lat)
    cx1, cy1 = geo_to_canvas_px(crop_max_lng, crop_min_lat)
    cx0, cy0 = max(0, int(cx0)), max(0, int(cy0))
    cx1, cy1 = min(canvas_w, int(cx1)), min(canvas_h, int(cy1))

    # Trim the crop box to exactly the output aspect ratio. Pixel space on
    # the tile canvas is uniform Web Mercator, so pixel aspect = ground
    # aspect — this guarantees the final resize is distortion-free.
    target_ar = out_w / out_h
    crop_w, crop_h = cx1 - cx0, cy1 - cy0
    if crop_w > 0 and crop_h > 0:
        if crop_w / crop_h > target_ar:      # too wide — trim width
            new_w = int(crop_h * target_ar)
            dx = (crop_w - new_w) // 2
            cx0 += dx; cx1 = cx0 + new_w
        else:                                # too tall — trim height
            new_h = int(crop_w / target_ar)
            dy = (crop_h - new_h) // 2
            cy0 += dy; cy1 = cy0 + new_h

    geo = (nw_lng, se_lat, se_lng, nw_lat)
    if cx1 > cx0 and cy1 > cy0:
        canvas = canvas.crop((cx0, cy0, cx1, cy1))
        # Geo extent of the cropped region
        w_full = canvas_w
        h_full = canvas_h
        west  = nw_lng + (cx0 / w_full) * (se_lng - nw_lng)
        east  = nw_lng + (cx1 / w_full) * (se_lng - nw_lng)
        north = nw_lat + (cy0 / h_full) * (se_lat - nw_lat)
        south = nw_lat + (cy1 / h_full) * (se_lat - nw_lat)
        geo = (west, south, east, north)

    canvas = canvas.resize((out_w, out_h), Image.LANCZOS)
    return canvas, geo


def draw_boundary(img, site_geom, geo_extent):
    w, h = img.size
    west, south, east, north = geo_extent
    lng_span = east  - west
    lat_span = north - south

    def geo_to_px(lng, lat):
        px = int((lng - west)  / lng_span * w)
        py = int((north - lat) / lat_span * h)
        return px, py

    draw = ImageDraw.Draw(img)
    lw   = max(3, w // 200)
    geoms = site_geom.geoms if hasattr(site_geom, 'geoms') else [site_geom]
    for poly in geoms:
        if not hasattr(poly, 'exterior'):
            continue
        pts = [geo_to_px(lng, lat) for lng, lat in poly.exterior.coords]
        if len(pts) >= 2:
            draw.line(pts + [pts[0]], fill=SITE_COLOR, width=lw)


# ── Word doc helpers ───────────────────────────────────────────────────────────
def get_or_add(el, tag):
    child = el.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        el.insert(0, child)
    return child


def tbl_border(tbl, val='single', sz=12, color='000000'):
    tblPr = get_or_add(tbl._tbl, 'w:tblPr')
    b = OxmlElement('w:tblBorders')
    for e in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{e}')
        el.set(qn('w:val'), val); el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0'); el.set(qn('w:color'), color)
        b.append(el)
    tblPr.append(b)


def no_border(cell):
    tcPr = get_or_add(cell._tc, 'w:tcPr')
    b = OxmlElement('w:tcBorders')
    for e in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{e}')
        el.set(qn('w:val'), 'none'); el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0'); el.set(qn('w:color'), 'auto')
        b.append(el)
    tcPr.append(b)


def cell_w(cell, emu):
    tcPr = get_or_add(cell._tc, 'w:tcPr')
    el   = OxmlElement('w:tcW')
    el.set(qn('w:w'), str(int(emu / 914400 * 1440)))
    el.set(qn('w:type'), 'dxa')
    tcPr.append(el)


def run(para, text, bold=False, italic=False, pt=11, font='Segoe UI', color=None):
    r = para.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.name = font; r.font.size = Pt(pt)
    if color:
        r.font.color.rgb = RGBColor(*color)
    return r


def build_doc_bytes(img_bytes, project_no, north_arrow_bytes):
    """Build the Word doc in memory and return bytes."""
    doc = Document()
    for sec in doc.sections:
        sec.page_width      = PAGE_W;  sec.page_height    = PAGE_H
        sec.left_margin     = MAR_L;   sec.right_margin   = MAR_R
        sec.top_margin      = MAR_TOP; sec.bottom_margin  = MAR_BOT
        sec.header_distance = Inches(0); sec.footer_distance = Inches(0)
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    cw    = PAGE_W - MAR_L - MAR_R
    cap_w = Inches(5.3)
    key_w = Inches(2.25)

    # Map image framed
    mt = doc.add_table(1, 1)
    tbl_border(mt)
    mt.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = get_or_add(mt._tbl, 'w:tblPr')
    tblW  = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(int(cw / 914400 * 1440)))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

    mc = mt.cell(0, 0)
    cell_w(mc, cw)
    tcPr  = get_or_add(mc._tc, 'w:tcPr')
    tcMar = OxmlElement('w:tcMar')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:w'), '0'); el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)

    mp = mc.paragraphs[0]
    mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mp.paragraph_format.space_before = Pt(0)
    mp.paragraph_format.space_after  = Pt(0)
    mp.add_run().add_picture(io.BytesIO(img_bytes), width=MAP_IMG_W, height=MAP_IMG_H)

    # Caption row
    ct = doc.add_table(1, 2)
    ct.alignment = WD_TABLE_ALIGNMENT.LEFT
    lc = ct.cell(0, 0); rc = ct.cell(0, 1)
    cell_w(lc, cap_w); cell_w(rc, key_w)
    no_border(lc); no_border(rc)

    lp1 = lc.paragraphs[0]
    lp1.paragraph_format.space_before = Pt(0)
    lp1.paragraph_format.space_after  = Pt(0)
    run(lp1, 'Figure 1:  Site Location Map', bold=True, pt=12, font='Segoe UI')

    lp2 = lc.add_paragraph()
    lp2.paragraph_format.space_before = Pt(0)
    lp2.paragraph_format.space_after  = Pt(0)
    run(lp2, f'Project No. {project_no}', bold=True, pt=10, font='Segoe UI')

    rp1 = rc.paragraphs[0]
    rp1.paragraph_format.space_before = Pt(0)
    rp1.paragraph_format.space_after  = Pt(0)
    run(rp1, 'KEY:', bold=True, pt=9)

    rp2 = rc.add_paragraph()
    rp2.paragraph_format.space_before = Pt(0)
    rp2.paragraph_format.space_after  = Pt(0)
    if north_arrow_bytes:
        rp2.add_run().add_picture(io.BytesIO(north_arrow_bytes), height=Inches(0.28))
        run(rp2, '  ', pt=9)
    sym = rp2.add_run('━━  ')
    sym.font.color.rgb = RGBColor(255, 0, 0); sym.font.size = Pt(12)
    run(rp2, 'Subject Property', pt=9)

    rp3 = rc.add_paragraph()
    rp3.paragraph_format.space_before = Pt(2)
    rp3.paragraph_format.space_after  = Pt(0)
    run(rp3, 'Drawing Not to Scale', italic=True, pt=8)

    # Footer
    ftr = doc.sections[0].footer
    for p in ftr.paragraphs:
        p._element.getparent().remove(p._element)
    ft  = ftr.add_table(1, 3, width=Inches(7.55))
    ft.alignment = WD_TABLE_ALIGNMENT.LEFT
    lc2 = ft.cell(0, 0); mc2 = ft.cell(0, 1); rc2 = ft.cell(0, 2)
    cell_w(lc2, Inches(4.0)); cell_w(mc2, Inches(1.5)); cell_w(rc2, Inches(2.05))
    for c in (lc2, mc2, rc2):
        no_border(c)

    lf  = lc2.paragraphs[0]
    r1  = lf.add_run('Figure 1:  Site Location Map')
    r1.bold = True; r1.font.name = 'Segoe UI'; r1.font.size = Pt(14); r1.font.all_caps = True
    lf2 = lc2.add_paragraph()
    r2  = lf2.add_run(f'Project No. {project_no}')
    r2.bold = True; r2.font.name = 'Segoe UI'; r2.font.size = Pt(10)

    rf = rc2.paragraphs[0]
    rf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rf.paragraph_format.space_before = Pt(6)
    cr = rf.add_run('Chenier Environmental Consulting, LLC')
    cr.font.name = 'Segoe UI'; cr.font.size = Pt(10)

    tP  = get_or_add(ft._tbl, 'w:tblPr')
    brd = OxmlElement('w:tblBorders')
    tp  = OxmlElement('w:top')
    tp.set(qn('w:val'), 'single'); tp.set(qn('w:sz'), '4'); tp.set(qn('w:color'), 'auto')
    brd.append(tp); tP.append(brd)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ── Streamlit UI ────────────────────────────────────────────────────────────────
st.set_page_config(page_title="Chenier — Site Location Map", page_icon="📄", layout="centered")

st.title("📄 Site Location Map Generator")
st.caption("Chenier Environmental Consulting, LLC — Phase I ESA Figure 1")

st.markdown("---")

uploaded = st.file_uploader("**1. Upload site boundary (KMZ or KML)**", type=['kmz', 'kml'])

col1, col2 = st.columns([2, 1])
with col1:
    project_no = st.text_input("**2. Project number**", placeholder="e.g. 26-014")
with col2:
    buffer_factor = st.slider("**Zoom out**", 0.3, 30.0, 1.0, 0.1,
                              help="Higher = more surrounding area shown")

generate = st.button("⚡ Generate Site Location Map", type="primary", use_container_width=True)

if generate:
    if not uploaded:
        st.error("Please upload a KMZ or KML file.")
    elif not project_no.strip():
        st.error("Please enter a project number.")
    else:
        try:
            with st.status("Generating map...", expanded=True) as status:
                st.write("Reading boundary file...")
                file_bytes = uploaded.read()
                site_geom, raw_bounds = parse_kmz_bytes(file_bytes, uploaded.name)

                min_lng, min_lat, max_lng, max_lat = raw_bounds
                lng_buf = max((max_lng - min_lng) * buffer_factor, MIN_BUFFER_DEG)
                lat_buf = max((max_lat - min_lat) * buffer_factor, MIN_BUFFER_DEG)
                buf_bounds = (min_lng - lng_buf, min_lat - lat_buf,
                              max_lng + lng_buf, max_lat + lat_buf)

                # Expand the shorter dimension in Web Mercator so the bbox
                # aspect ratio matches the output image — prevents stretch.
                def _to_merc(lng, lat):
                    x = lng * 20037508.34 / 180.0
                    y = math.log(math.tan((90 + lat) * math.pi / 360.0)) / (math.pi / 180.0)
                    return x, y * 20037508.34 / 180.0

                def _to_lnglat(x, y):
                    lng = x / 20037508.34 * 180.0
                    lat = math.degrees(2 * math.atan(
                        math.exp(y / 20037508.34 * math.pi)) - math.pi / 2)
                    return lng, lat

                mx0, my0 = _to_merc(buf_bounds[0], buf_bounds[1])
                mx1, my1 = _to_merc(buf_bounds[2], buf_bounds[3])
                bw, bh = mx1 - mx0, my1 - my0
                target_ar = IMG_W_PX / IMG_H_PX
                cx, cy = (mx0 + mx1) / 2, (my0 + my1) / 2
                if bw / bh < target_ar:    # too tall — widen
                    bw = bh * target_ar
                else:                      # too wide — heighten
                    bh = bw / target_ar
                lng0, lat0 = _to_lnglat(cx - bw / 2, cy - bh / 2)
                lng1, lat1 = _to_lnglat(cx + bw / 2, cy + bh / 2)
                buf_bounds = (lng0, lat0, lng1, lat1)

                zoom = choose_zoom(buf_bounds, IMG_W_PX, IMG_H_PX)
                st.write(f"Fetching OpenStreetMap tiles (zoom {zoom})...")

                prog = st.progress(0.0)
                canvas, geo_extent = build_osm_canvas(
                    buf_bounds, zoom, IMG_W_PX, IMG_H_PX,
                    progress=lambda f: prog.progress(f))

                st.write("Drawing site boundary...")
                draw_boundary(canvas, site_geom, geo_extent)

                img_buf = io.BytesIO()
                canvas.save(img_buf, 'JPEG', quality=JPEG_QUALITY, optimize=True)
                img_bytes = img_buf.getvalue()

                st.write("Building Word document...")
                north_arrow_bytes = None
                try:
                    with open(NORTH_ARROW_PATH, 'rb') as f:
                        north_arrow_bytes = f.read()
                except FileNotFoundError:
                    st.write("(north_arrow.jpeg not found — north arrow omitted)")

                safe_proj = project_no.strip().replace('/', '_').replace('\\', '_')
                docx_bytes = build_doc_bytes(img_bytes, project_no.strip(), north_arrow_bytes)

                status.update(label="Done!", state="complete", expanded=False)

            st.success("Site Location Map generated.")
            st.image(canvas, caption="Preview", use_container_width=True)

            fname = f"{safe_proj}_Fig_1_Site_Location_Map.docx"
            st.download_button(
                "⬇ Download Word Document",
                data=docx_bytes,
                file_name=fname,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                use_container_width=True,
            )

        except Exception as e:
            st.error(f"Error: {e}")
            import traceback
            st.code(traceback.format_exc())

st.markdown("---")
st.caption("Map data © OpenStreetMap contributors")

